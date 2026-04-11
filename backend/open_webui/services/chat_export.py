"""Chat export service — PDF and Word generation."""

import logging
from io import BytesIO
from pathlib import Path

from jinja2 import Environment, FileSystemLoader, select_autoescape

from open_webui.utils.chat_export import prepare_export_messages

log = logging.getLogger(__name__)

TEMPLATE_DIR = Path(__file__).parent.parent / 'templates'


def _render_html(title: str, messages: list[dict]) -> str:
    """Render chat as styled HTML using Jinja2 template."""
    prepared_messages, sources = prepare_export_messages(messages)

    env = Environment(
        loader=FileSystemLoader(str(TEMPLATE_DIR)),
        autoescape=select_autoescape(default_for_string=True, default=True),
    )
    template = env.get_template('chat_export.html')

    return template.render(
        title=title,
        messages=prepared_messages,
        sources=sources,
    )


def generate_pdf(title: str, messages: list[dict]) -> bytes:
    """Generate PDF from chat messages."""
    from weasyprint import HTML

    html_string = _render_html(title, messages)
    pdf_bytes = HTML(string=html_string).write_pdf()
    return pdf_bytes


def generate_docx(title: str, messages: list[dict]) -> bytes:
    """Generate DOCX from chat messages using python-docx directly."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from htmldocx import HtmlToDocx

    prepared_messages, sources = prepare_export_messages(messages)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title_para = doc.add_heading(title, level=1)

    # Messages
    parser = HtmlToDocx()
    for msg in prepared_messages:
        # Role header
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(msg['role'].upper())
        role_run.bold = True
        role_run.font.size = Pt(10)
        role_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        role_para.space_after = Pt(4)

        # Content — use htmldocx for the message body HTML only
        html_content = msg.get('html_content', '')
        if html_content.strip():
            parser.add_html_to_document(f'<div>{html_content}</div>', doc)

        # Add spacing after message
        spacer = doc.add_paragraph()
        spacer.space_before = Pt(0)
        spacer.space_after = Pt(8)

    # Sources section
    if sources:
        # Divider line
        doc.add_paragraph('').space_after = Pt(4)

        sources_heading = doc.add_heading('Bronnen', level=2)

        for source in sources:
            source_para = doc.add_paragraph()
            # Number in blue bold
            num_run = source_para.add_run(f'[{source.index}] ')
            num_run.bold = True
            num_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
            num_run.font.size = Pt(10)

            # Name
            name_run = source_para.add_run(source.name)
            name_run.font.size = Pt(10)

            # URL if present
            if source.url and source.url != source.name:
                url_run = source_para.add_run(f' — {source.url}')
                url_run.font.size = Pt(10)
                url_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

            source_para.space_after = Pt(2)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
